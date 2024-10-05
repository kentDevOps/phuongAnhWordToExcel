import os
from docx import Document
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import Alignment,PatternFill
from log import readIni
import re
from unidecode import unidecode
#import xlsxwriter
#import xlwings as xw
#import win32com.client as win32

# Hàm đọc file Word và lọc các dòng chứa cụm [CÔNG CHỨC *]
# Hàm ghi dữ liệu vào cột A của Sheet1 và nối vào dòng cuối cùng
def contains_roman_numerals(text):
    # Regular expression to match Roman numerals
    roman_pattern = r'\b[IVXLCDM]+\b'
    return re.search(roman_pattern, text) is not None
def count_lines_in_word(word_path):
    doc = Document(word_path)
    total_lines = 0
    
    # Đếm số dòng dựa trên dấu xuống dòng trong mỗi đoạn văn
    for para in doc.paragraphs:
        # Mỗi đoạn văn có thể chứa nhiều dòng nếu có ký tự '\n'
        total_lines += para.text.count('\n') + 1 if para.text.strip() else 0

    return total_lines
def append_to_excel(cauHoi,daA,daB,daC,daD,kq,giaiThich, excel_path,report_path,dem):
    # Mở workbook và chọn sheet 'Sheet1'
    wb = load_workbook(excel_path)
    ws = wb['s1']
    
    # Tìm dòng trống cuối cùng trong cột A
    last_row = find_last_row(ws,column='A') +1
    
    # Ghi dữ liệu vào cột A, bắt đầu từ dòng cuối cùng
    for idx, line in enumerate(cauHoi, start=last_row):
        ws[f'A{idx}'] = line
     # Ghi dữ liệu vào cột C, bắt đầu từ dòng cuối cùng
    for idx, line in enumerate(daA, start=last_row):
        ws[f'C{idx}'] = line
     # Ghi dữ liệu vào cột D, bắt đầu từ dòng cuối cùng
    for idx, line in enumerate(daB, start=last_row):
        ws[f'D{idx}'] = line
     # Ghi dữ liệu vào cột E, bắt đầu từ dòng cuối cùng
    for idx, line in enumerate(daC, start=last_row):
        ws[f'E{idx}'] = line
     # Ghi dữ liệu vào cột F, bắt đầu từ dòng cuối cùng
    for idx, line in enumerate(daD, start=last_row):
        ws[f'F{idx}'] = line
     # Ghi dữ liệu vào cột G, bắt đầu từ dòng cuối cùng
    for idx, line in enumerate(kq, start=last_row):
        ws[f'G{idx}'] = line
            # Ghi dữ liệu vào cột H, bắt đầu từ dòng cuối cùng
    for idx, line in enumerate(giaiThich, start=last_row):
        ws[f'H{idx}'] = line                         
    # Lưu lại file Excel
    last_row = find_last_row(ws,column='A') +1
    for row in range(1, ws.max_row + 1):  # Bắt đầu từ hàng 1 đến hàng cuối
        cell = ws[f'{'H'}{row}']
        cell.alignment = Alignment(wrap_text=True)  
        if cell.value:
            num_lines = len(str(cell.value).split('\n'))  # Đếm số dòng trong ô
            ws.row_dimensions[row].height = 35 * num_lines  # Điều chỉnh hệ số 15 cho từng dòng  
    ws.title = os.path.basename(report_path)
    text_without_accents = unidecode( ws.title)
    newName = os.path.join(os.path.dirname(report_path),text_without_accents) 
    wb.save(newName)
    print(f"Kết quả đã được thêm vào {newName} tại Folder Report...")
'''def boido(report_path):
# Mở ứng dụng Excel
    # Bước 1: Mở file Excel hiện có bằng openpyxl
    input_file = report_path  # Đường dẫn tới file Excel hiện có
    output_file = report_path  # File Excel mới sẽ được tạo
    workbook = xlsxwriter.Workbook(report_path)
    ws = workbook.shee
    # Bước 4: Đóng file mới
    wb_new.close() '''
def filter_congchuc_lines(word_path):
    doc = Document(word_path)
    total = 0
    for para in doc.paragraphs:
        # Mỗi đoạn văn có thể chứa nhiều dòng nếu có ký tự '\n'
        total += para.text.count('\n') + 1 if para.text.strip() else 0    
    questions = []
    strDk = readIni("cauHoi","strListCau")
    listDk = strDk.split(',')
    gt_index = -1
    flag = 0
    for index,para in enumerate(doc.paragraphs):
        line = para.text
        print(line)
        
        if "Giải thích" in line:
            flag = 1
        elif "Dịch nghĩa" in line:
            flag = 0
        if "[CÔNG CHỨC" in line:
            #line = line.split("\n")[0]
            question = line.split("[CÔNG CHỨC")[1]  # Lấy phần sau cụm "[CÔNG CHỨC"
            question = "[CÔNG CHỨC" + question.split("\n")[0]  # Thêm lại cụm "[CÔNG CHỨC"
            questions.append(question.strip())  
        if "[CÔNG CHỨC" not in line and 'VOCABULARY' not in line and  'EXERCISES' not in line:
            gt_index = index
            if  gt_index < total:
                if "Dịch nghĩa" in doc.paragraphs[gt_index +1].text :
                    pass
                elif "Dịch nghĩa" not in doc.paragraphs[gt_index +1].text and flag==0 and "[CÔNG CHỨC" not in doc.paragraphs[gt_index +1].text:
                    for item in listDk:
                        if item in line:   
                            question = line.split(item)[1]  # Lấy phần sau cụm "[CÔNG CHỨC"
                            question =item + question.split("\n")[0]  # Thêm lại cụm "[CÔNG CHỨC"  
                            questions.append(question.strip())
                            break;                
    return questions
def filter_DapAnA_lines(word_path,dapAn):
    doc = Document(word_path)
    questions = []
    i = 0
    for para in doc.paragraphs:
        line = para.text
        if dapAn in line:
            i+=1
            #line = line.split("\n")[0]
            if 'B.' in line and 'C.' not in line and ":" not in line:
                question = line.split("B.")[0]
                questions.append(question.strip())
            elif 'A.' in line and 'B.' in line and 'C.' in line and ":" not in line:
                start = line.find('A.') + len('A.')
                end = line.find('B.')
                question = line[start:end].strip()
                questions.append(question.strip()) 
            elif 'B.' not in line and 'C.' not in line and ":" not in line:
                question = line.split(dapAn)[1]  # Lấy phần sau cụm "[CÔNG CHỨC"
                question = question.split("\n")[0]  # Thêm lại cụm "[CÔNG CHỨC"
                questions.append(question.strip())            
    return questions  
def filter_DapAnB_lines(word_path,dapAn):
    doc = Document(word_path)
    questions = []
    i = 0
    for para in doc.paragraphs:
        line = para.text
        if dapAn in line:
            i+=1
            #line = line.split("\n")[0]
            if dapAn in line and 'C.' not in line and ":" not in line:
                question = line.split("B.")[1]
                questions.append(question.strip())
            elif 'A.' in line and 'B.' in line and 'C.' in line and ":" not in line:
                start = line.find('B.') + len('B.')
                end = line.find('C.')
                question = line[start:end].strip()
                questions.append(question.strip())           
    return questions  
def filter_DapAnC_lines(word_path,dapAn):
    doc = Document(word_path)
    questions = []
    i = 0
    for para in doc.paragraphs:
        line = para.text
        if dapAn in line:
            i+=1
            #line = line.split("\n")[0]
            if dapAn in line and 'D.' not in line and ":" not in line:
                question = line.split(dapAn)[1]
                questions.append(question.strip())
            elif 'A.' in line and 'B.' in line and 'C.' in line and ":" not in line:
                start = line.find(dapAn) + len(dapAn)
                end = line.find('D.')
                question = line[start:end].strip()
                questions.append(question.strip())           
    return questions    
def filter_DapAnD_lines(word_path,dapAn):
    doc = Document(word_path)
    questions = []
    i = 0
    for para in doc.paragraphs:
        line = para.text
        if dapAn in line:
            i+=1
            #line = line.split("\n")[0]
            if dapAn in line and 'C.' not in line and ":" not in line:
                question = line.split(dapAn)[1]
                questions.append(question.strip())
            elif 'A.' in line and 'B.' in line and 'C.' in line and ":" not in line:
                question = line.split(dapAn)[1]
                questions.append(question.strip())         
    return questions  
def find_last_row(ws, column="A"):
    last_row = ws.max_row
    while last_row > 1 and ws[f"{column}{last_row}"].value is None:
        last_row -= 1
    return last_row

#append_to_excel(rs,'rp.xlsx')
def extract_correct_answers(file_word):
    doc = Document(file_word)
    correct_answers = []
    
    for para in doc.paragraphs:
        # Chỉ xử lý các đoạn văn có chứa các đáp án A., B., C., D.
        if any(option in para.text for option in ['A.', 'B.', 'C.', 'D.']):
            answer = ""
            for run in para.runs:
                # Kiểm tra nếu đoạn được in đậm và nằm trong cụm đáp án
                if run.bold and ('A.' in run.text or 'B.' in run.text or 'C.' in run.text or 'D.' in run.text):
                    answer += run.text.strip()  # Lấy đoạn in đậm
            if answer:
                correct_answers.append(answer)
    #chuyen doi sang dap an dang so
    arrQuidoi = []
    for item in correct_answers:
        dapAn = str(item).split('.')[0].strip()
        if dapAn == "A":
            arrQuidoi.append(1)
        elif dapAn == "B":
            arrQuidoi.append(2)
        elif dapAn == "C":
            arrQuidoi.append(3)
        elif dapAn == "D":
            arrQuidoi.append(4)
    return arrQuidoi
def find_explanation_index(file_word):
    doc = Document(file_word)   
    gt_index = -1  # Khởi tạo với giá trị không hợp lệ
    arrGt = []
    flag = 0
    strGt = ""
    for index, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        temp = ""
        
        # Kiểm tra nếu dòng chứa từ "Giải thích"
        if "Giải thích" in text:
            flag = 1
            gt_index = index
            if "Dịch nghĩa" in doc.paragraphs[gt_index +1].text:
                t1 = 'Gt:' + doc.paragraphs[gt_index].text.split(':')[1]
                strGt = strGt + '\n' + t1
            else:
                strGt = strGt + '\n' + doc.paragraphs[gt_index].text 
        elif "A."  in text and "B."  in text and "C."  in text and "D."  in text and "Dịch nghĩa" not in text:      
            flag = 0      
        elif "A."  in text and "B."  in text and "C."  in text and "D."  in text and "Dịch nghĩa"  in text: 
            temp = text.split("A.")[0]
            strGt = strGt + '\n' + temp
            flag = 0   
        elif "Dịch nghĩa" in text and "[CÔNG CHỨC "  not in text:
            gt_index = index
            strGt = strGt + '\n' + doc.paragraphs[gt_index].text 
            flag = 0   
        elif "Tạm dịch" in text:
            gt_index = index
            strGt = strGt + '\n' + doc.paragraphs[gt_index].text 
            flag = 0                                      
        elif "Giải thích" not in text and flag == 1 and "[CÔNG CHỨC " not in text and '\xa0' not in text :
            gt_index = index
            strGt = strGt + '\n' + doc.paragraphs[gt_index].text    
            #flag = 0
        elif "[CÔNG CHỨC "  in text and "Dịch nghĩa" in text:
            temp = text.split("[CÔNG CHỨC ")[0]
            strGt = strGt + '\n' + temp
            flag = 0
        elif "[CÔNG CHỨC " in text and "Dịch nghĩa" not in text and flag == 1 and "A."  in text: 
            flag = 0           
        elif "[CÔNG CHỨC " in text and "Dịch nghĩa" not in text and flag == 1 and "A." not in text:
            flag = 0
            temp = text.split("[CÔNG CHỨC ")[0]
            if len(temp.strip())>4:
                strGt = strGt + '\n' + temp
        elif "[CÔNG CHỨC " not in text and "Giải thích" not in text and "Dịch nghĩa" not in text and flag == 0:  
            flag = 0      
        elif "[CÔNG CHỨC " not in text and "Giải thích" not in text and "Dịch nghĩa" not in text and flag == 1 and '\xa0' in text:
            flag = 0     
              
    arrStrGt = strGt.split("Giải thích") 
    arrStrGt.pop(0)  
    arrStrGt = ['Giải thích:' + item for item in arrStrGt]
    for i, explanation in enumerate(arrStrGt, 1):
        print(f"Giải thích {i}:\n{explanation}\n")            
    return arrStrGt


# Lấy phần giải thích từ file
'''explanations = find_explanation_index(file_word)
#print(explanations)
# In ra các phần giải thích đã trích xuất
for i, explanation in enumerate(explanations, 1):
    print(f"Giải thích {i}:\n{explanation}\n")
correct_answers = extract_correct_answers(file_word)
dA = filter_DapAnD_lines(file_word,'D.')'''